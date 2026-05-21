"""
Generate the overall design report DOCX and ER diagram from docs/overall_design.md.

The script intentionally keeps the source of truth in Markdown, then renders the
course-facing Word files so the text docs and DOCX deliverables stay aligned.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


BASE_DIR = Path(__file__).resolve().parent
SOURCE_MD = BASE_DIR / "overall_design.md"
ER_IMAGE = BASE_DIR / "er_diagram.png"
OUTPUT_DOCX = [
    BASE_DIR / "总体设计报告.docx",
    BASE_DIR / "电子商务平台管理系统总体设计报告.docx",
    BASE_DIR / "电子商务平台管理系统总体设计报告_含ER图.docx",
]


TABLE_GROUPS = [
    (
        "用户与权限",
        "#e0f2fe",
        [
            "t_user",
            "t_address",
            "t_points_log",
            "t_operation_log",
        ],
    ),
    (
        "商品与库存",
        "#dcfce7",
        [
            "t_category",
            "t_product",
            "t_tag",
            "t_product_tag",
            "t_spec_template",
            "t_spec_value",
            "t_product_sku",
        ],
    ),
    (
        "交易与营销",
        "#fef3c7",
        [
            "t_shopping_cart",
            "t_order",
            "t_order_item",
            "t_coupon",
            "t_user_coupon",
            "t_seckill_session",
            "t_seckill_product",
        ],
    ),
    (
        "售后与消息",
        "#fce7f3",
        [
            "t_refund",
            "t_notification",
            "t_notification_read",
        ],
    ),
]

RELATIONS = [
    ("t_user", "t_order", "1:N"),
    ("t_user", "t_shopping_cart", "1:N"),
    ("t_user", "t_user_coupon", "1:N"),
    ("t_user", "t_refund", "1:N"),
    ("t_product", "t_product_sku", "1:N"),
    ("t_product", "t_order_item", "1:N"),
    ("t_product", "t_product_tag", "1:N"),
    ("t_tag", "t_product_tag", "1:N"),
    ("t_coupon", "t_user_coupon", "1:N"),
    ("t_order", "t_order_item", "1:N"),
    ("t_order", "t_refund", "1:0..1"),
    ("t_seckill_session", "t_seckill_product", "1:N"),
    ("t_notification", "t_notification_read", "1:N"),
]


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def hex_to_rgb(value: str) -> tuple[int, int, int]:
    value = value.lstrip("#")
    return tuple(int(value[i : i + 2], 16) for i in (0, 2, 4))


def draw_round_rect(draw: ImageDraw.ImageDraw, box, fill, outline="#334155", radius=18):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=2)


def draw_centered_text(draw: ImageDraw.ImageDraw, box, text: str, font, fill="#0f172a"):
    left, top, right, bottom = box
    bbox = draw.textbbox((0, 0), text, font=font)
    x = left + (right - left - (bbox[2] - bbox[0])) / 2
    y = top + (bottom - top - (bbox[3] - bbox[1])) / 2
    draw.text((x, y), text, font=font, fill=fill)


def generate_er_diagram() -> None:
    width, height = 2200, 1420
    image = Image.new("RGB", (width, height), "#f8fafc")
    draw = ImageDraw.Draw(image)
    title_font = load_font(46, bold=True)
    group_font = load_font(30, bold=True)
    table_font = load_font(24)
    relation_font = load_font(20)

    draw.text((70, 42), "电子商务平台管理系统 ER 关系总览", font=title_font, fill="#0f172a")
    draw.text(
        (72, 105),
        "22 张主要业务表，覆盖用户、商品、交易、营销、售后和消息通知",
        font=load_font(24),
        fill="#475569",
    )

    group_boxes = [
        (70, 180, 610, 980),
        (640, 180, 1180, 980),
        (1210, 180, 1750, 980),
        (1780, 180, 2130, 980),
    ]

    for (title, color, tables), box in zip(TABLE_GROUPS, group_boxes):
        draw_round_rect(draw, box, fill="#ffffff", outline="#cbd5e1", radius=28)
        draw.rounded_rectangle(
            (box[0] + 22, box[1] + 22, box[2] - 22, box[1] + 86),
            radius=18,
            fill=hex_to_rgb(color),
            outline="#cbd5e1",
            width=1,
        )
        draw_centered_text(draw, (box[0] + 22, box[1] + 22, box[2] - 22, box[1] + 86), title, group_font)
        y = box[1] + 122
        for table in tables:
            table_box = (box[0] + 44, y, box[2] - 44, y + 70)
            draw_round_rect(draw, table_box, fill=hex_to_rgb(color), outline="#94a3b8", radius=14)
            draw_centered_text(draw, table_box, table, table_font)
            y += 92

    relation_box = (70, 1030, 2130, 1340)
    draw_round_rect(draw, relation_box, fill="#ffffff", outline="#cbd5e1", radius=28)
    draw.text((relation_box[0] + 32, relation_box[1] + 26), "关键关系", font=group_font, fill="#0f172a")
    columns = 3
    col_width = (relation_box[2] - relation_box[0] - 80) // columns
    start_x = relation_box[0] + 34
    start_y = relation_box[1] + 88
    for idx, (source, target, label) in enumerate(RELATIONS):
        col = idx % columns
        row = idx // columns
        x = start_x + col * col_width
        y = start_y + row * 48
        relation = f"{source}  {label}  {target}"
        draw.text((x, y), relation, font=relation_font, fill="#334155")

    image.save(ER_IMAGE)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_run_font(run, size: int | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)

    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    for style_name in ["Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
        style = document.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_inline_markdown(paragraph, text: str, size: int | None = None) -> None:
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            if size:
                run.font.size = Pt(size)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size)


def parse_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells):
            continue
        rows.append(cells)
    return rows


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    table.autofit = True
    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            add_inline_markdown(paragraph, row[c_idx] if c_idx < len(row) else "", size=9)
            if r_idx == 0:
                set_cell_shading(cell, "E2E8F0")
                for run in paragraph.runs:
                    run.bold = True
    document.add_paragraph()


def add_code_block(document: Document, code_lines: list[str]) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    for index, line in enumerate(code_lines):
        if index:
            paragraph.add_run("\n")
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(8.5)
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F1F5F9")
    p_pr.append(shading)


def add_title_page(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(12)
    run = paragraph.add_run("电子商务平台管理系统总体设计报告")
    set_run_font(run, size=24, bold=True, color="0F172A")

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline_markdown(subtitle, "基于当前 dev 分支代码、SQL 脚本与项目文档修订", size=12)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline_markdown(meta, "修订时间：2026-05-21", size=11)
    document.add_page_break()


def render_markdown(document: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []
    table_lines: list[str] = []

    def flush_table() -> None:
        nonlocal table_lines
        if table_lines:
            add_table(document, parse_table(table_lines))
            table_lines = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            flush_table()
            if in_code:
                add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if line.strip().startswith("|") and line.strip().endswith("|"):
            table_lines.append(line)
            i += 1
            continue
        flush_table()

        stripped = line.strip()
        if not stripped:
            document.add_paragraph()
        elif stripped.startswith("![") and ER_IMAGE.exists():
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(str(ER_IMAGE), width=Inches(6.7))
        elif stripped.startswith("# "):
            heading = document.add_heading(stripped[2:], level=1)
            heading.paragraph_format.space_before = Pt(8)
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            document.add_heading(stripped[4:], level=3)
        elif stripped.startswith("#### "):
            document.add_heading(stripped[5:], level=4)
        elif stripped.startswith(">"):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.5)
            run = paragraph.add_run(stripped.lstrip("> ").strip())
            set_run_font(run, size=10, color="475569")
            run.italic = True
        elif stripped.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_markdown(paragraph, stripped[2:])
        elif re.match(r"^\d+\. ", stripped):
            paragraph = document.add_paragraph()
            add_inline_markdown(paragraph, re.sub(r"^(\d+)\. ", r"\1、", stripped))
        else:
            paragraph = document.add_paragraph()
            add_inline_markdown(paragraph, stripped)

        i += 1

    flush_table()
    if code_lines:
        add_code_block(document, code_lines)


def build_docx(output_path: Path) -> None:
    markdown = SOURCE_MD.read_text(encoding="utf-8")
    document = Document()
    style_document(document)
    add_title_page(document)
    render_markdown(document, markdown)

    # Keep the ER page readable if Word opens the picture section on narrow pages.
    for section in document.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)

    document.save(output_path)


def main() -> None:
    generate_er_diagram()
    for output_path in OUTPUT_DOCX:
        build_docx(output_path)
        print(f"generated: {output_path}")


if __name__ == "__main__":
    main()
